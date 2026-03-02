"""
Document text extraction module.
Supports PDF, DOCX, and plain text files.
"""

import io
from typing import Optional
from docx import Document as DocxDocument
from pypdf import PdfReader


class DocumentParser:
    """Parser for extracting text from various document formats."""

    def __init__(self, file_content: bytes, filename: str):
        """
        Initialize the parser with file content.

        Args:
            file_content: Raw bytes of the uploaded file
            filename: Name of the uploaded file (used to determine format)
        """
        self.file_content = file_content
        self.filename = filename
        self.file_ext = self._get_file_extension()

    def _get_file_extension(self) -> str:
        """Extract and return lowercase file extension."""
        if '.' in self.filename:
            return self.filename.rsplit('.', 1)[-1].lower()
        return ''

    def extract_text(self) -> str:
        """
        Extract text from the document based on file type.

        Returns:
            Extracted text as a string

        Raises:
            ValueError: If file format is not supported
        """
        if self.file_ext == 'pdf':
            return self._extract_from_pdf()
        elif self.file_ext == 'docx':
            return self._extract_from_docx()
        elif self.file_ext in ('txt', 'text'):
            return self._extract_from_text()
        else:
            raise ValueError(
                f"Unsupported file format: '.{self.file_ext}'. "
                "Supported formats: PDF, DOCX, TXT"
            )

    def _extract_from_pdf(self) -> str:
        """Extract text from a PDF file."""
        try:
            pdf_file = io.BytesIO(self.file_content)
            reader = PdfReader(pdf_file)
            text_parts = []

            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text_parts.append(extracted)

            return '\n'.join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to parse PDF file: {str(e)}")

    def _extract_from_docx(self) -> str:
        """Extract text from a DOCX file."""
        try:
            docx_file = io.BytesIO(self.file_content)
            doc = DocxDocument(docx_file)

            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(' | '.join(row_text))

            return '\n'.join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX file: {str(e)}")

    def _extract_from_text(self) -> str:
        """Extract text from a plain text file."""
        try:
            return self.file_content.decode('utf-8')
        except UnicodeDecodeError:
            # Try with different encodings
            for encoding in ['latin-1', 'iso-8859-1', 'cp1252']:
                try:
                    return self.file_content.decode(encoding)
                except (UnicodeDecodeError, LookupError):
                    continue
            raise ValueError("Failed to decode text file. Please ensure it's UTF-8 encoded.")